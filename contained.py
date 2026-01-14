import webview, math, random, asyncio
from docx import Document


class API:

  # suggested attempt to resolve error by making sure "Document" was being run in a synchronous method (did not work)
  def load_doc(self, path: str):
    return Document(path)

  # Method takes a given 'template file' overwrites the marked {code} sections with the value of the code
  async def write_code(self, codes):
    output = []

    # print(codes)

    i = 0
    file_count = 1
    # Loops through provided list of `codes`
    while i < len(codes):
      # !!!!!!! BREAKS HERE !!!!!!!!
      # Opens template file
      # Attempts to open template file within synchronous and asynchronous method (neither method works)
      # doc = Document('./templates/Template_1.docx')
      doc = await asyncio.to_thread(self.load_doc, './templates/Template_1.docx')

      # inserts codes into template file
      for p in doc.paragraphs:
        if '{code1}' in p.text:
          p.text = p.text.replace('{code1}', codes[i]['specification']['code'])
          output.append(f'Saved File: `./to_print/custom_file_{i}.docx` with code: `{i}`')
        try:
          # Saves filled Template File
          doc.save(f'./to_print/custom_file_{file_count}.docx')
        except Exception as e:
          print(e)
    
      i += 1
      file_count += 1

    # print(output)
    return output

  # Generates a random test code [no issues]
  async def generate_code(self, copies, amount):
    gen_name = str(random.randint(1000, 9999)) + "_Name"
    gen_code = str(random.randint(1000, 9999)) + "_Code"
    amount = float(amount)

    payload = {
      "specification": {
        "name": gen_name,
        "code": gen_code,
        "amount": amount
      }
    }

    return payload


  # Called by .html file.
  # Required to run. Unclear why. Other methods all failed
  def submit_generation_request(self, request):
    return asyncio.run(self._submit_generation_request(request))
  # async def submit_generation_request(self, request):
  #   return await self._submit_generation_request(request)
  # def submit_generation_request(self, request):
  #   future = asyncio.run_coroutine_threadsafe(
  #       self._submit_generation_request(request),
  #       self.loop  # stored event loop
  #   )
  #   return future.result()

  # Method takes user given information, cleans the data, generates the requested number of codes and submits `generated_codes` to `write_code` to insert the codes into a file
  async def _submit_generation_request(self, request):
    # ensures `copy_count` is treated as a float and sets it to the absolute value
    copy_count = abs(float(request[0]))

    # case for decimal value
    copy_count = math.ceil(copy_count)
    
    # case for odd number selected
    if copy_count % 2 == 1:
      copy_count += 1

    # tracks list of generated codes that will be inserted into the template file
    generated_codes = []
    # generates requested codes equal to user specified `copy_count`
    for i in range(int(copy_count)):
      # generates random stand in codes
      generated_codes.append(await self.generate_code(request[0], request[1]))

    # sends `generated_codes` off to be inserted into the template_file
    output = await self.write_code(generated_codes)
    print(output)

    return f"Success!"


# Launches UI Elements [no issues]
async def main():
  api = API()
  window = webview.create_window(
    "User Input",
    "contained_input.html",
    js_api=api,
    width=500,
    height=600,
  )
  webview.start()


if __name__ == "__main__":
  # Allows webview to be run inside of an async method
  asyncio.run(main())
